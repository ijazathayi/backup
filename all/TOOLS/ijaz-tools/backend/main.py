"""
IJAZ Tools — Unified Backend
FastAPI server powering all tools: Network Monitor, Nmap Studio, IP Scanner, Training Scripts, AI Chat.
"""

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from fastapi.background import BackgroundTasks
import asyncio
import json
import os
import re
import shutil
import socket
import subprocess
import platform
import tempfile
import time
import ipaddress
import urllib.request
import queue as _queue
from concurrent.futures import ThreadPoolExecutor
from typing import AsyncGenerator

# Load .env if present (no external dependency needed)
_env_path = os.path.join(os.path.dirname(__file__), ".env")
if os.path.isfile(_env_path):
    with open(_env_path) as _f:
        for _line in _f:
            _line = _line.strip()
            if _line and not _line.startswith("#") and "=" in _line:
                _k, _, _v = _line.partition("=")
                os.environ.setdefault(_k.strip(), _v.strip())

# ─────────────────────────────────────────────
# App setup
# ─────────────────────────────────────────────

app = FastAPI(title="IJAZ Tools API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

executor = ThreadPoolExecutor(max_workers=20)

# ─────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────

def get_local_ip() -> str:
    """Return the machine's LAN IP address."""
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(("8.8.8.8", 80))
        return s.getsockname()[0]
    except Exception:
        return "127.0.0.1"
    finally:
        s.close()


def run_cmd(cmd: str, timeout: int = 10) -> str:
    """Run a shell command and return combined stdout+stderr."""
    result = subprocess.run(
        cmd, shell=True, capture_output=True, text=True,
        timeout=timeout, errors="replace"
    )
    return result.stdout + result.stderr


# ─────────────────────────────────────────────
# General API
# ─────────────────────────────────────────────

@app.get("/api/status")
def status():
    return {"status": "online", "version": "1.0.0"}


@app.get("/api/myip")
def my_ip():
    return {"ip": get_local_ip()}


# ─────────────────────────────────────────────
# Network Monitor
# ─────────────────────────────────────────────

_subnet_cache: dict = {"data": None, "ts": 0}
_ports_cache: dict = {"data": 0, "ts": 0}
_bw_state: dict = {"recv": 0, "sent": 0, "ts": time.time()}


def _get_subnets() -> tuple[str, str]:
    now = time.time()
    if _subnet_cache["data"] and now - _subnet_cache["ts"] < 30:
        return _subnet_cache["data"]
    try:
        out = run_cmd("ipconfig", timeout=5)
        wifi = hotspot = ""
        adapter = ""
        for raw in out.splitlines():
            line = raw.strip()
            if raw and not raw[0].isspace() and line.endswith(":"):
                adapter = line.lower()
            elif "IPv4 Address" in line:
                ip = line.split(":")[-1].strip()
                subnet = ".".join(ip.split(".")[:3]) + "."
                if "wi-fi" in adapter:
                    wifi = subnet
                elif "local area connection*" in adapter:
                    hotspot = subnet
        _subnet_cache["data"] = (wifi, hotspot)
        _subnet_cache["ts"] = now
        return wifi, hotspot
    except Exception:
        return "", ""


def get_devices(mode: str = "wifi") -> list[dict]:
    try:
        wifi, hotspot = _get_subnets()
        target = hotspot if mode == "hotspot" and hotspot else wifi
        if not target:
            return []
        out = run_cmd("arp -a", timeout=5)
        devices = []
        for line in out.splitlines():
            if "dynamic" in line.lower():
                parts = line.split()
                if len(parts) >= 2 and parts[0].startswith(target):
                    devices.append({"ip": parts[0], "mac": parts[1], "name": "Unknown"})
        return devices
    except Exception:
        return []


def get_open_ports() -> int:
    now = time.time()
    if now - _ports_cache["ts"] < 10:
        return _ports_cache["data"]
    try:
        out = run_cmd("netstat -an", timeout=5)
        count = sum(1 for l in out.splitlines() if "LISTENING" in l)
        _ports_cache["data"] = count
        _ports_cache["ts"] = now
        return count
    except Exception:
        return 0


def get_bandwidth() -> tuple[str, int, int]:
    try:
        out = run_cmd("netstat -e", timeout=5)
        for line in out.splitlines():
            if line.startswith("Bytes"):
                parts = line.split()
                recv, sent = int(parts[1]), int(parts[2])
                now = time.time()
                elapsed = max(0.1, now - _bw_state["ts"])
                recv_bps = max(0, (recv - _bw_state["recv"]) / elapsed)
                sent_bps = max(0, (sent - _bw_state["sent"]) / elapsed)
                _bw_state.update({"recv": recv, "sent": sent, "ts": now})
                max_b = 50 * 1024 * 1024
                total_mb = (recv_bps + sent_bps) / (1024 * 1024)
                return (
                    f"{total_mb:.2f} MB/s",
                    min(100, int(recv_bps / max_b * 100)),
                    min(100, int(sent_bps / max_b * 100)),
                )
    except Exception:
        pass
    return "0.00 MB/s", 0, 0


def get_ping_stats() -> tuple[int, int]:
    try:
        out = run_cmd("ping -n 1 -w 1000 8.8.8.8", timeout=4)
        if re.search(r"100% loss|unreachable|could not find host", out, re.I):
            return 0, 0
        m = re.search(r"time[=<](\d+)ms", out, re.I)
        ms = int(m.group(1)) if m else 0
        return 100, max(0, min(100, 100 - ms // 5))
    except Exception:
        return 0, 0


@app.get("/api/network/devices")
def api_devices(mode: str = "wifi"):
    return get_devices(mode)


@app.get("/api/network/stats")
def api_stats(mode: str = "wifi"):
    devices = get_devices(mode)
    bw, dl, ul = get_bandwidth()
    ph, ps = get_ping_stats()
    return {
        "online": len(devices),
        "ports": get_open_ports(),
        "bandwidth": bw,
        "traffic": {"download": dl, "upload": ul, "packetHealth": ph, "pingStability": ps},
    }


@app.get("/api/network/scan/{ip}")
def api_scan_device(ip: str):
    if not re.match(r"^\d{1,3}(\.\d{1,3}){3}$", ip):
        return JSONResponse({"error": "Invalid IP"}, status_code=400)
    lines = ["=== PING ==="]
    try:
        lines.append(run_cmd(f"ping -n 4 {ip}", timeout=15).strip())
    except Exception as e:
        lines.append(f"Ping failed: {e}")
    lines.append("\n=== NMAP SCAN ===")
    try:
        lines.append(run_cmd(f"nmap -A -T4 -F {ip}", timeout=90).strip())
    except Exception:
        try:
            lines.append(run_cmd(f"nmap -sV -F {ip}", timeout=90).strip())
        except Exception as e:
            lines.append(f"Nmap failed: {e}")
    return {"output": "\n".join(lines), "target": ip}


@app.get("/api/network/scan-host")
def api_scan_host(mode: str = "wifi"):
    try:
        wifi, hotspot = _get_subnets()
        target = hotspot if mode == "hotspot" and hotspot else wifi
        if not target:
            return JSONResponse({"error": "Could not determine host IP"}, status_code=400)
        out = run_cmd("ipconfig", timeout=5)
        my_ip = None
        for line in out.splitlines():
            if "IPv4 Address" in line:
                ip = line.split(":")[-1].strip()
                if ip.startswith(target):
                    my_ip = ip
                    break
        if not my_ip:
            return JSONResponse({"error": "Could not find host IP"}, status_code=400)
        return {"output": run_cmd(f"nmap -F -sV {my_ip}", timeout=90), "target": my_ip}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


# ─────────────────────────────────────────────
# Network Monitor — WebSocket live feed
# ─────────────────────────────────────────────

network_clients: set[WebSocket] = set()
current_mode = "wifi"


async def push_network_update():
    if not network_clients:
        return
    devices = get_devices(current_mode)
    bw, dl, ul = get_bandwidth()
    ph, ps = get_ping_stats()
    payload = json.dumps({
        "type": "update",
        "devices": devices,
        "stats": {
            "online": len(devices),
            "ports": get_open_ports(),
            "bandwidth": bw,
            "traffic": {"download": dl, "upload": ul, "packetHealth": ph, "pingStability": ps},
        },
    })
    dead = set()
    for ws in network_clients:
        try:
            await ws.send_text(payload)
        except Exception:
            dead.add(ws)
    network_clients.difference_update(dead)


@app.websocket("/ws/network")
async def ws_network(websocket: WebSocket):
    global current_mode
    await websocket.accept()
    network_clients.add(websocket)
    try:
        await push_network_update()
        while True:
            try:
                raw = await asyncio.wait_for(websocket.receive_text(), timeout=3.0)
                msg = json.loads(raw)
                if msg.get("type") == "set_mode":
                    current_mode = msg.get("mode", "wifi")
                await push_network_update()
            except asyncio.TimeoutError:
                await push_network_update()
    except (WebSocketDisconnect, Exception):
        pass
    finally:
        network_clients.discard(websocket)


# ─────────────────────────────────────────────
# Nmap Studio — WebSocket streaming scan
# ─────────────────────────────────────────────

COMMON_PORTS = {
    20: "FTP Data", 21: "FTP Control", 22: "SSH", 23: "Telnet", 25: "SMTP",
    53: "DNS", 80: "HTTP", 110: "POP3", 143: "IMAP", 443: "HTTPS",
    445: "SMB", 1433: "MSSQL", 3306: "MySQL", 3389: "RDP",
    5432: "PostgreSQL", 8080: "HTTP Alt", 8443: "HTTPS Alt",
}


def enrich_line(line: str) -> str:
    m = re.match(r"^(\d+)/(tcp|udp)\s+open\s+", line.strip())
    if m:
        port = int(m.group(1))
        desc = COMMON_PORTS.get(port, "Open port — service unspecified")
        return line.rstrip() + f"  ← {desc}\n"
    return line


async def stream_nmap(target: str, flags: str) -> AsyncGenerator[str, None]:
    cmd = ["nmap"] + flags.split() + [target]
    proc = subprocess.Popen(
        cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
        text=True, bufsize=1, errors="replace"
    )
    loop = asyncio.get_running_loop()
    while True:
        try:
            line = await asyncio.wait_for(
                loop.run_in_executor(executor, proc.stdout.readline), timeout=3.0
            )
            if not line:
                break
            yield enrich_line(line)
        except asyncio.TimeoutError:
            yield "__KEEPALIVE__"
    proc.wait()


@app.websocket("/ws/scan")
async def ws_scan(websocket: WebSocket):
    await websocket.accept()
    try:
        raw = await websocket.receive_text()
        params = json.loads(raw)
        target = params.get("target", "").strip()
        flags = params.get("flags", "-sV -T4").strip()

        if not target:
            await websocket.send_text(json.dumps({"type": "error", "data": "Target is required"}))
            return

        async for line in stream_nmap(target, flags):
            if line == "__KEEPALIVE__":
                await websocket.send_text(json.dumps({"type": "keepalive"}))
            else:
                await websocket.send_text(json.dumps({"type": "output", "data": line}))

        await websocket.send_text(json.dumps({"type": "done"}))
    except (WebSocketDisconnect, Exception) as e:
        try:
            await websocket.send_text(json.dumps({"type": "error", "data": str(e)}))
        except Exception:
            pass
    finally:
        try:
            await websocket.close()
        except Exception:
            pass


# ─────────────────────────────────────────────
# IP Scanner — ping sweep
# ─────────────────────────────────────────────

def ping_host(ip: str) -> bool:
    param = "-n" if platform.system() == "Windows" else "-c"
    wait = ["-w", "300"] if platform.system() == "Windows" else ["-W", "1"]
    try:
        r = subprocess.run(
            ["ping", param, "1"] + wait + [ip],
            capture_output=True, text=True, timeout=2, errors="replace"
        )
        return "TTL" in r.stdout or "ttl=" in r.stdout.lower()
    except Exception:
        return False


@app.websocket("/ws/ipscan")
async def ws_ipscan(websocket: WebSocket):
    await websocket.accept()
    try:
        raw = await websocket.receive_text()
        params = json.loads(raw)
        subnet_input = params.get("subnet", "").strip()

        if not subnet_input:
            local_ip = get_local_ip()
            subnet_input = f"{local_ip}/24"

        try:
            network = ipaddress.ip_network(subnet_input, strict=False)
        except ValueError:
            await websocket.send_text(json.dumps({"type": "error", "data": "Invalid subnet"}))
            return

        hosts = list(network.hosts())
        total = len(hosts)
        await websocket.send_text(json.dumps({"type": "start", "total": total, "subnet": str(network)}))

        loop = asyncio.get_running_loop()
        found = 0

        async def scan_one(ip_obj):
            nonlocal found
            ip = str(ip_obj)
            alive = await loop.run_in_executor(executor, ping_host, ip)
            if alive:
                found += 1
                await websocket.send_text(json.dumps({"type": "found", "ip": ip}))

        batch = 50
        for i in range(0, total, batch):
            chunk = hosts[i:i + batch]
            await asyncio.gather(*[scan_one(ip) for ip in chunk])
            progress = min(100, int((i + batch) / total * 100))
            await websocket.send_text(json.dumps({"type": "progress", "pct": progress}))

        await websocket.send_text(json.dumps({"type": "done", "found": found}))
    except (WebSocketDisconnect, Exception) as e:
        try:
            await websocket.send_text(json.dumps({"type": "error", "data": str(e)}))
        except Exception:
            pass
    finally:
        try:
            await websocket.close()
        except Exception:
            pass


# ─────────────────────────────────────────────
# Python Scripts — WebSocket streaming runner
# ─────────────────────────────────────────────

SCRIPTS_DIR = os.path.join(os.path.dirname(__file__), "scripts")

ALLOWED_SCRIPTS = {
    "pynmap":           {"file": "pynmap.py",          "needs_target": True},
    "network_scanner":  {"file": "network_scanner.py", "needs_target": False},
    "port_scanner":     {"file": "port_scanner.py",    "needs_target": True},
    "service_detector": {"file": "service_detector.py","needs_target": True},
    "packet_sniffer":   {"file": "packet_sniffer.py",  "needs_target": False},
    "maclookup":        {"file": "maclookup.py",       "needs_target": False},
    "net_toolkit":      {"file": "net_toolkit.py",     "needs_target": True},
}

_ARG_SAFE = re.compile(r"^[a-zA-Z0-9.\-_/:, ]+$")


def _build_cmd(script_id: str, target: str, flags: str):
    """Validate inputs and return the command list, or raise ValueError."""
    if script_id not in ALLOWED_SCRIPTS:
        raise ValueError("Script not in allowlist")

    meta = ALLOWED_SCRIPTS[script_id]
    if meta["needs_target"] and not target:
        raise ValueError("This script requires a target IP or subnet")

    script_path = os.path.join(SCRIPTS_DIR, meta["file"])
    if not os.path.isfile(script_path):
        raise FileNotFoundError(f"Script file not found: {meta['file']}")

    if target and not _ARG_SAFE.match(target):
        raise ValueError("Invalid characters in target")
    if flags and not _ARG_SAFE.match(flags):
        raise ValueError("Invalid characters in flags")

    cmd = ["python", "-u", script_path]   # -u = unbuffered stdout
    if target:
        cmd.append(target)
    if flags:
        cmd += flags.split()
    return cmd


@app.websocket("/ws/script")
async def ws_script(websocket: WebSocket):
    """
    Stream a Python script's output line-by-line.
    Client sends: {"script": "pynmap", "target": "...", "flags": "..."}
    Server sends: {"type": "output"|"done"|"error", "data": "..."}
    """
    await websocket.accept()
    proc = None
    try:
        raw    = await websocket.receive_text()
        params = json.loads(raw)

        script_id = params.get("script", "").strip()
        target    = str(params.get("target", "")).strip()
        flags     = str(params.get("flags",  "")).strip()

        try:
            cmd = _build_cmd(script_id, target, flags)
        except (ValueError, FileNotFoundError) as e:
            await websocket.send_text(json.dumps({"type": "error", "data": str(e)}))
            return

        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,          # line-buffered
            errors="replace",
        )

        loop = asyncio.get_running_loop()

        while True:
            try:
                line = await asyncio.wait_for(
                    loop.run_in_executor(executor, proc.stdout.readline),
                    timeout=2.0,
                )
            except asyncio.TimeoutError:
                # Send a heartbeat so the UI knows we're still alive
                await websocket.send_text(json.dumps({"type": "heartbeat"}))
                # Check if process already finished
                if proc.poll() is not None:
                    break
                continue

            if not line:
                break

            await websocket.send_text(json.dumps({"type": "output", "data": line}))

        proc.wait()
        exit_code = proc.returncode
        await websocket.send_text(json.dumps({
            "type": "done",
            "data": f"\n[Process exited with code {exit_code}]",
        }))

    except WebSocketDisconnect:
        pass
    except Exception as e:
        try:
            await websocket.send_text(json.dumps({"type": "error", "data": str(e)}))
        except Exception:
            pass
    finally:
        if proc and proc.poll() is None:
            proc.kill()
        try:
            await websocket.close()
        except Exception:
            pass


# ─────────────────────────────────────────────
# OCR — PDF upload → OCR → Word download
# ─────────────────────────────────────────────

OCR_REPO = os.path.join(os.path.dirname(__file__), "..", "..", "OCRmyPDF")


def _check_deps() -> list[str]:
    """Return list of missing dependencies."""
    missing = []
    try:
        import ocrmypdf  # noqa: F401
    except ImportError:
        missing.append("ocrmypdf")
    try:
        import pdfminer  # noqa: F401
    except ImportError:
        missing.append("pdfminer.six")
    try:
        import docx  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    return missing


@app.get("/api/ocr/status")
def ocr_status():
    missing = _check_deps()
    return {"ready": len(missing) == 0, "missing": missing}


@app.post("/api/ocr/convert")
async def ocr_convert(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    language: str = Form("eng"),
    deskew: bool = Form(False),
):
    """Upload a PDF → run OCRmyPDF → extract text → return .docx"""
    missing = _check_deps()
    if missing:
        return JSONResponse(
            {"error": f"Missing packages: {', '.join(missing)}. Run: pip install {' '.join(missing)}"},
            status_code=503,
        )

    if not file.filename.lower().endswith(".pdf"):
        return JSONResponse({"error": "Only PDF files are supported"}, status_code=400)

    import ocrmypdf
    from pdfminer.high_level import extract_text as pdf_extract_text
    from docx import Document
    from docx.shared import Pt

    tmp_dir     = tempfile.mkdtemp(prefix="ijaz_ocr_")
    input_path  = os.path.join(tmp_dir, "input.pdf")
    ocr_path    = os.path.join(tmp_dir, "ocr.pdf")
    output_path = os.path.join(tmp_dir, "output.docx")
    stem        = os.path.splitext(file.filename)[0]
    orig_name   = file.filename

    # Save upload
    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    loop = asyncio.get_running_loop()

    def _build_docx_from_text(text: str, path: str, title: str):
        doc = Document()
        doc.add_heading(title, level=1)
        for block in text.split("\n\n"):
            block = block.strip()
            if block:
                para = doc.add_paragraph(block.replace("\n", " "))
                para.style.font.size = Pt(11)
        doc.save(path)

    try:
        # Try OCR path first
        def _run_ocr():
            ocrmypdf.ocr(
                input_path, ocr_path,
                language=language,
                deskew=deskew,
                progress_bar=False,
                skip_text=True,
            )

        try:
            await loop.run_in_executor(executor, _run_ocr)
            source_pdf = ocr_path
        except ocrmypdf.exceptions.PriorOcrFoundError:
            # PDF already has a text layer — use it directly
            source_pdf = input_path

        text = await loop.run_in_executor(executor, pdf_extract_text, source_pdf)
        await loop.run_in_executor(executor, _build_docx_from_text, text, output_path, stem)

        # Clean up tmp dir after response is sent
        background_tasks.add_task(shutil.rmtree, tmp_dir, True)

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{stem}_ocr.docx",
        )

    except Exception as e:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return JSONResponse({"error": str(e)}, status_code=500)


# ─────────────────────────────────────────────
# AI Assistant — streaming chat via OpenAI
# ─────────────────────────────────────────────

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
OPENAI_MODEL   = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

AI_SYSTEM_PROMPT = """You are a knowledgeable network security assistant built into IJAZ Tools — 
a local network operations console. You help with:
- Nmap scan interpretation and flag recommendations
- Network troubleshooting and diagnostics
- Port and service explanations
- Python networking scripts
- General cybersecurity concepts and best practices

Be concise, practical, and technical. Format code blocks with triple backticks."""


@app.websocket("/ws/chat")
async def ws_chat(websocket: WebSocket):
    """
    Streaming OpenAI chat.
    Client sends: {"messages": [{"role": "user"|"assistant", "content": "..."}]}
    Server sends: {"type": "token", "data": "..."} per chunk, then {"type": "done"}
    """
    await websocket.accept()
    try:
        raw  = await websocket.receive_text()
        body = json.loads(raw)
        messages = body.get("messages", [])

        # Accept key from client (set in browser localStorage) or fall back to server env
        api_key = str(body.get("api_key", "")).strip() or OPENAI_API_KEY

        if not api_key:
            await websocket.send_text(json.dumps({
                "type": "error",
                "data": "No API key found. Click '🔑 Set API key' in the chat header and paste your OpenAI key."
            }))
            return

        if not messages:
            await websocket.send_text(json.dumps({"type": "error", "data": "No messages provided"}))
            return

        payload = json.dumps({
            "model": OPENAI_MODEL,
            "stream": True,
            "messages": [{"role": "system", "content": AI_SYSTEM_PROMPT}] + messages,
        }).encode()

        req = urllib.request.Request(
            "https://api.openai.com/v1/chat/completions",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            },
            method="POST",
        )

        loop = asyncio.get_running_loop()
        token_queue: _queue.Queue = _queue.Queue()

        def _stream_to_queue():
            try:
                with urllib.request.urlopen(req, timeout=60) as resp:
                    for raw_line in resp:
                        line = raw_line.decode("utf-8").strip()
                        if not line.startswith("data:"):
                            continue
                        data = line[5:].strip()
                        if data == "[DONE]":
                            break
                        try:
                            chunk = json.loads(data)
                            token = chunk["choices"][0]["delta"].get("content", "")
                            if token:
                                token_queue.put(("token", token))
                        except Exception:
                            pass
            except Exception as e:
                token_queue.put(("error", str(e)))
            finally:
                token_queue.put(("done", None))

        future = loop.run_in_executor(executor, _stream_to_queue)

        while True:
            try:
                kind, value = token_queue.get_nowait()
            except _queue.Empty:
                if future.done():
                    while not token_queue.empty():
                        kind, value = token_queue.get_nowait()
                        if kind == "token":
                            await websocket.send_text(json.dumps({"type": "token", "data": value}))
                        elif kind == "error":
                            await websocket.send_text(json.dumps({"type": "error", "data": value}))
                        elif kind == "done":
                            await websocket.send_text(json.dumps({"type": "done"}))
                    break
                await asyncio.sleep(0.02)
                continue

            if kind == "token":
                await websocket.send_text(json.dumps({"type": "token", "data": value}))
            elif kind == "error":
                await websocket.send_text(json.dumps({"type": "error", "data": value}))
                break
            elif kind == "done":
                await websocket.send_text(json.dumps({"type": "done"}))
                break

    except WebSocketDisconnect:
        pass
    except Exception as e:
        try:
            await websocket.send_text(json.dumps({"type": "error", "data": str(e)}))
        except Exception:
            pass
    finally:
        try:
            await websocket.close()
        except Exception:
            pass


# ─────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
